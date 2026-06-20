"""Generacion de memoria tecnica en Word para la aplicacion hidrologica."""
from __future__ import annotations
from io import BytesIO
from typing import Any, Dict, Optional
import os
import tempfile

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _set_cell_shading(cell, fill="D9EAF7"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)


def add_df_table(doc: Document, df: pd.DataFrame, title: str = None, max_rows: int = 200):
    if title:
        doc.add_heading(title, level=2)
    if df is None or df.empty:
        doc.add_paragraph("Sin datos disponibles.")
        return
    dfx = df.copy().head(max_rows)
    dfx = dfx.fillna("—")
    table = doc.add_table(rows=1, cols=len(dfx.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(dfx.columns):
        hdr[i].text = str(col)
        _set_cell_shading(hdr[i], "BDD7EE")
    for _, row in dfx.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, (float, int)):
                try:
                    fv = float(val)
                    if abs(fv - round(fv)) < 1e-9 and abs(fv) >= 20:
                        txt = f"{int(round(fv)):,}".replace(",", ".")
                    else:
                        txt = f"{fv:,.3f}".replace(",", "X").replace(".", ",").replace("X", ".")
                except Exception:
                    txt = str(val)
            else:
                txt = str(val)
            cells[i].text = txt
    _format_table(table)


def add_formula_paragraph(doc: Document, formula: str, desc: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula)
    run.bold = True
    run.font.size = Pt(11)
    if desc:
        doc.add_paragraph(desc)


def save_profile_plot(profile_df: pd.DataFrame, path: str):
    fig, ax = plt.subplots(figsize=(7, 3.2), dpi=160)
    if profile_df is not None and not profile_df.empty:
        ax.plot(profile_df["dist_m"] / 1000.0, profile_df["cota_m"])
    ax.set_xlabel("Distancia acumulada (km)")
    ax.set_ylabel("Cota estimada (m s.n.m.)")
    ax.set_title("Perfil longitudinal del cauce principal")
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_results_plot(results_df: pd.DataFrame, path: str):
    fig, ax = plt.subplots(figsize=(7, 3.4), dpi=160)
    if results_df is not None and not results_df.empty:
        df = results_df.dropna(subset=["Q_m3s"])
        if not df.empty:
            piv = df.pivot_table(index="T_anios", columns="metodo", values="Q_m3s", aggfunc="first")
            for col in piv.columns:
                ax.plot(piv.index, piv[col], marker="o", label=str(col))
            ax.legend(fontsize=7, loc="best")
    ax.set_xlabel("Periodo de retorno (años)")
    ax.set_ylabel("Caudal máximo estimado (m³/s)")
    ax.set_title("Comparación de caudales por metodología")
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_debris_plot(debris_df: pd.DataFrame, path: str):
    fig, ax = plt.subplots(figsize=(7, 3.4), dpi=160)
    if debris_df is not None and not debris_df.empty and "QD_m3s" in debris_df.columns:
        df = debris_df.dropna(subset=["QD_m3s"])
        if not df.empty:
            piv = df.pivot_table(index="T_anios", columns="escenario_Cv", values="QD_m3s", aggfunc="max")
            for col in piv.columns:
                ax.plot(piv.index, piv[col], marker="o", label=str(col))
            ax.legend(fontsize=7, loc="best")
    ax.set_xlabel("Periodo de retorno (años)")
    ax.set_ylabel("Caudal detrítico QD (m³/s)")
    ax.set_title("Caudal detrítico por escenario de concentración volumétrica")
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def generate_memoria_docx(
    output_path: str,
    project_name: str,
    objective: str,
    metrics: Dict[str, Any],
    contour_metrics: Dict[str, Any],
    profile_summary: Dict[str, Any],
    methods_df: pd.DataFrame,
    results_df: pd.DataFrame,
    profile_df: pd.DataFrame,
    debris_df: Optional[pd.DataFrame] = None,
    volume_df: Optional[pd.DataFrame] = None,
    empirical_peak_df: Optional[pd.DataFrame] = None,
    susceptibility: Optional[Dict[str, Any]] = None,
    tc_methods_df: Optional[pd.DataFrame] = None,
    design_flow_df: Optional[pd.DataFrame] = None,
    pmax24_freq_df: Optional[pd.DataFrame] = None,
    rain_crosscheck_df: Optional[pd.DataFrame] = None,
    basin_image_path: Optional[str] = None,
    surface_image_path: Optional[str] = None,
    notes: str = "",
):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(9)
    for h in ['Heading 1', 'Heading 2', 'Heading 3']:
        styles[h].font.name = 'Arial'
        styles[h].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_heading("Memoria de cálculo hidrológico y morfométrico", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(project_name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_paragraph(f"Objetivo del cálculo: {objective}")

    doc.add_heading("1. Resumen ejecutivo", level=1)
    A = metrics.get("area_km2")
    P = metrics.get("perimeter_km")
    L = profile_summary.get("L_km")
    H = profile_summary.get("H_m")
    S = profile_summary.get("S_pct")
    Tc = profile_summary.get("Tc_kirpich_min")
    tc_adopt = profile_summary.get("Tc_adoptado_min", Tc)
    doc.add_paragraph(
        f"La cuenca analizada presenta un área aproximada de {A:.3f} km², perímetro de {P:.3f} km "
        f"y longitud hidráulica adoptada/preliminar de {L:.3f} km. El desnivel hidráulico estimado es de "
        f"{H:.2f} m, equivalente a una pendiente media de {S:.2f} %. El tiempo de concentración adoptado "
        f"automáticamente es {tc_adopt:.2f} minutos, obtenido a partir de una comparación multimétodo cuando existe información suficiente."
    )
    doc.add_paragraph(
        "Los resultados hidrológicos se presentan por metodología y periodo de retorno. Los métodos parametrizados "
        "DGA-AC y Verni-King modificado requieren que el usuario ingrese coeficientes oficiales/regionales para su "
        "uso como valor de diseño."
    )

    if basin_image_path and os.path.exists(basin_image_path):
        doc.add_heading("2. Imagen de referencia de la cuenca", level=1)
        doc.add_picture(basin_image_path, width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if surface_image_path and os.path.exists(surface_image_path):
        doc.add_picture(surface_image_path, width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3. Parámetros morfométricos", level=1)
    morph_df = pd.DataFrame([
        ["Área", metrics.get("area_km2"), "km²"],
        ["Área", metrics.get("area_ha"), "ha"],
        ["Perímetro", metrics.get("perimeter_km"), "km"],
        ["Longitud geométrica máxima", metrics.get("max_geom_length_km"), "km"],
        ["Ancho medio", metrics.get("mean_width_km"), "km"],
        ["Índice de compacidad Kc", metrics.get("compactness_kc"), "-"],
        ["Factor de forma", metrics.get("form_factor"), "-"],
        ["Relación de elongación", metrics.get("elongation_ratio"), "-"],
        ["Centroide lon", metrics.get("centroid_lon"), "°"],
        ["Centroide lat", metrics.get("centroid_lat"), "°"],
        ["EPSG UTM adoptado", metrics.get("epsg_utm"), "-"],
    ], columns=["Parámetro", "Valor", "Unidad"])
    add_df_table(doc, morph_df)

    doc.add_heading("4. Curvas de nivel y altimetría", level=1)
    cont_df = pd.DataFrame([
        ["Número de curvas", contour_metrics.get("n_contours"), "-"],
        ["Número de niveles", contour_metrics.get("n_levels"), "-"],
        ["Cota mínima detectada", contour_metrics.get("z_min"), "m"],
        ["Cota máxima detectada", contour_metrics.get("z_max"), "m"],
        ["Equidistancia estimada", contour_metrics.get("equidistance_m"), "m"],
        ["Vértices de curvas", contour_metrics.get("n_vertices"), "-"],
    ], columns=["Parámetro", "Valor", "Unidad"])
    add_df_table(doc, cont_df)

    doc.add_heading("5. Perfil longitudinal y tiempo de concentración", level=1)
    prof_sum_df = pd.DataFrame([
        ["Longitud hidráulica", profile_summary.get("L_km"), "km"],
        ["Cota inicial", profile_summary.get("z_start_m"), "m"],
        ["Cota final", profile_summary.get("z_end_m"), "m"],
        ["Desnivel", profile_summary.get("H_m"), "m"],
        ["Pendiente media", profile_summary.get("S_pct"), "%"],
        ["Pendiente Mociornita", profile_summary.get("S_mociornita_pct"), "%"],
        ["Tc Kirpich", profile_summary.get("Tc_kirpich_min"), "min"],
        ["Tc Giandotti", profile_summary.get("Tc_giandotti_min"), "min"],
        ["Tc adoptado automático", profile_summary.get("Tc_adoptado_min"), "min"],
    ], columns=["Parámetro", "Valor", "Unidad"])
    add_df_table(doc, prof_sum_df)
    if tc_methods_df is not None and not tc_methods_df.empty:
        add_df_table(doc, tc_methods_df, "5.1 Comparación de métodos de tiempo de concentración")

    if profile_df is not None and not profile_df.empty and profile_df.get("cota_m", pd.Series(dtype=float)).notna().any():
        with tempfile.TemporaryDirectory() as td:
            prof_plot = os.path.join(td, "perfil.png")
            save_profile_plot(profile_df, prof_plot)
            if os.path.exists(prof_plot):
                doc.add_picture(prof_plot, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("No se incorporó perfil longitudinal gráfico porque no se definió un cauce principal válido. La longitud y pendiente usadas son preliminares.")

    doc.add_heading("6. Fórmulas principales", level=1)
    add_formula_paragraph(doc, "P_t^T = P_{D,10} · CD_t · CF_T · K · CA", "Estimación de lluvia de diseño desde lluvia diaria de 10 años, coeficientes de duración/frecuencia, corrección K y abatimiento espacial CA.")
    add_formula_paragraph(doc, "I_t^T = P_t^T / t", "Intensidad de lluvia para duración t, usada en el método racional cuando no existe IDF local completa.")
    add_formula_paragraph(doc, "Q = 0,278 · C · I · A", "Método racional: Q en m³/s, C adimensional, I en mm/h y A en km².")
    add_formula_paragraph(doc, "Tc = 0,0195 · L^0,77 · S^-0,385", "Kirpich: Tc en minutos, L en metros y S en m/m.")
    add_formula_paragraph(doc, "S = Δh/A · (l0/2 + Σli + ln/2)", "Pendiente media de cuenca según Mociornita: Δh en m, A en m² y longitudes de curvas en m.")
    add_formula_paragraph(doc, "Tc = (0,87 · L³ / H)^0,385", "Método California: Tc en horas, L en km y H en m.")
    add_formula_paragraph(doc, "Tc = L/(v·3600)", "Métodos US Navy y Velocidad Texas: L en m, v según pendiente/cobertura y Tc en horas.")
    add_formula_paragraph(doc, "Pe = (P - 0,2S)^2 / (P + 0,8S)", "SCS-CN: precipitación efectiva para P mayor que la abstracción inicial.")
    add_formula_paragraph(doc, "Qp = 0,208 · A · Pe / Tp", "Hidrograma unitario SCS simplificado: A en km², Pe en mm y Tp en horas.")
    add_formula_paragraph(doc, "QVK = k · P24^m · A^n", "Verni-King modificado parametrizado: los coeficientes deben ser ingresados desde fuente oficial/regional.")
    add_formula_paragraph(doc, "QDGA = qref · A^alpha · FT · Kinst", "DGA-AC parametrizado: requiere parámetros regionales, factor de frecuencia y conversión a instantáneo.")
    add_formula_paragraph(doc, "QD = QL + QS", "Caudal detrítico total como suma del gasto líquido y gasto sólido.")
    add_formula_paragraph(doc, "QD = QL / (1 - Cv)", "Conversión por concentración volumétrica de sedimentos Cv.")
    add_formula_paragraph(doc, "QS = QD - QL", "Gasto sólido equivalente asociado al caudal detrítico.")
    add_formula_paragraph(doc, "Cv = [ρ/(σ-ρ)] · [tan(θ)/(tan(φ)-tan(θ))]", "Estimación simplificada de concentración volumétrica según Takahashi, usando pendiente de cauce y propiedades del sedimento.")
    add_formula_paragraph(doc, "Qp = a · M^b", "Relaciones empíricas de caudal detrítico máximo en función del volumen de sedimento transportado.")

    add_df_table(doc, methods_df, "7. Diagnóstico de metodologías")
    if pmax24_freq_df is not None and not pmax24_freq_df.empty:
        add_df_table(doc, pmax24_freq_df, "7.1 Análisis de frecuencia pluviométrica histórico")
    if rain_crosscheck_df is not None and not rain_crosscheck_df.empty:
        add_df_table(doc, rain_crosscheck_df, "7.2 Comparación P10D de isoyeta versus estaciones históricas")
    add_df_table(doc, results_df, "8. Resultados hidrológicos")
    if design_flow_df is not None and not design_flow_df.empty:
        add_df_table(doc, design_flow_df, "8.1 Caudal recomendado automático por período de retorno")

    with tempfile.TemporaryDirectory() as td:
        res_plot = os.path.join(td, "resultados.png")
        save_results_plot(results_df, res_plot)
        if os.path.exists(res_plot):
            doc.add_picture(res_plot, width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("9. Análisis sedimentológico y caudal detrítico", level=1)
    if susceptibility:
        doc.add_paragraph(
            f"Susceptibilidad aluvional/detrítica preliminar: {susceptibility.get('nivel', 'No evaluada')} "
            f"(puntaje {susceptibility.get('puntaje', '-')}). Fundamentos: {susceptibility.get('fundamentos', '')}."
        )
    doc.add_paragraph(
        "El caudal líquido calculado por los métodos hidrológicos se transforma en gasto sólido y caudal detrítico "
        "mediante escenarios de concentración volumétrica. Estos resultados son de prediseño y deben validarse con "
        "granulometría, disponibilidad real de sedimentos, secciones hidráulicas, evidencias de terreno y/o modelación especializada."
    )
    if debris_df is not None and not debris_df.empty:
        add_df_table(doc, debris_df, "9.1 Conversión Q líquido – Q sólido – Q detrítico")
        with tempfile.TemporaryDirectory() as td:
            deb_plot = os.path.join(td, "detritico.png")
            save_debris_plot(debris_df, deb_plot)
            if os.path.exists(deb_plot):
                doc.add_picture(deb_plot, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("No se calcularon caudales detríticos porque el módulo fue desactivado o faltaron caudales líquidos/Cv válidos.")
    if volume_df is not None and not volume_df.empty:
        add_df_table(doc, volume_df, "9.2 Volumen sedimentario estimado")
    if empirical_peak_df is not None and not empirical_peak_df.empty:
        add_df_table(doc, empirical_peak_df, "9.3 Relaciones empíricas Qp=f(M)")

    doc.add_heading("10. Análisis y advertencias técnicas", level=1)
    doc.add_paragraph(
        "La selección metodológica debe considerar tamaño de cuenca, disponibilidad de información pluviométrica y fluviométrica, "
        "objetivo de diseño, régimen hidrológico y existencia de obras de regulación. Cuando la cuenca se ubica cerca del límite "
        "superior del método racional, se recomienda contrastar con métodos regionales y métodos de hidrograma."
    )
    doc.add_paragraph(
        "Los métodos DGA-AC y Verni-King se implementan como módulos parametrizados con presets regionales editables. "
        "La aplicación además compara tiempos de concentración, aplica pendiente de Mociornita cuando hay curvas suficientes, "
        "contrasta isoyetas con series históricas y genera un caudal recomendado automático que no reemplaza el criterio del especialista."
    )
    if notes:
        doc.add_paragraph(notes)

    doc.add_heading("11. Conclusión", level=1)
    doc.add_paragraph(
        "La memoria entrega una base trazable para revisar morfometría, pendiente, tiempo de concentración, lluvia de diseño y caudales "
        "por metodología. El caudal final de diseño debe ser adoptado por el especialista, dejando constancia del criterio utilizado y "
        "de la fuente de los parámetros regionales o pluviométricos."
    )

    doc.save(output_path)
    return output_path


def generate_memoria_bytes(*args, **kwargs) -> bytes:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = tmp.name
    generate_memoria_docx(path, *args, **kwargs)
    with open(path, "rb") as f:
        data = f.read()
    try:
        os.remove(path)
    except OSError:
        pass
    return data
